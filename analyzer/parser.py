import PyPDF2
import docx
import re
from typing import Dict, List, Optional
import os
import logging

logger = logging.getLogger(__name__)

class CVParser:
    def __init__(self):
        # ---------------- CONTACT PATTERNS ----------------
        self.contact_patterns = [
            r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}\b',
            r'(?:\+\d{1,3}\s?)?\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}',
            r'\+\d{1,3}[-.\s]?\d{1,4}[-.\s]?\d{1,4}[-.\s]?\d{1,9}',
            r'\b\d{10,15}\b'
        ]

        # ---------------- SECTION KEYWORDS ----------------
        self.experience_keywords = [
            'experience', 'work experience', 'employment',
            'professional experience', 'career', 'work history'
        ]
        self.education_keywords = [
            'education', 'academic', 'degree', 'university',
            'college', 'qualification', 'certification'
        ]
        self.skills_keywords = [
            'skills', 'technical skills', 'competencies',
            'technologies', 'tools', 'expertise'
        ]

        self.stop_sections = {
            'experience': ['education', 'skills', 'projects', 'certifications'],
            'education': ['experience', 'skills', 'projects'],
            'skills': ['experience', 'education', 'projects']
        }

    # --------------------------------------------------
    # TEXT EXTRACTION
    # --------------------------------------------------

    def extract_text_from_pdf(self, file_path: str) -> str:
        if not os.path.exists(file_path):
            return f"Error: file not found - {file_path}"

        try:
            with open(file_path, "rb") as f:
                reader = PyPDF2.PdfReader(f)
                text = ""

                if reader.is_encrypted:
                    try:
                        reader.decrypt("")
                    except Exception:
                        return "Error: PDF is password protected"

                for page in reader.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"

                return text.strip() or "No readable text found in PDF"
        except Exception as e:
            logger.error(f"PDF read error: {e}")
            return f"Error reading PDF: {e}"

    def extract_text_from_docx(self, file_path: str) -> str:
        if not os.path.exists(file_path):
            return f"Error: file not found - {file_path}"

        try:
            doc = docx.Document(file_path)
            text = ""

            for p in doc.paragraphs:
                if p.text.strip():
                    text += p.text + "\n"

            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text += cell.text + "\n"

            return text.strip() or "No readable text found in DOCX"
        except Exception as e:
            logger.error(f"DOCX read error: {e}")
            return f"Error reading DOCX: {e}"

    def extract_text_from_txt(self, file_path: str) -> str:
        if not os.path.exists(file_path):
            return f"Error: file not found - {file_path}"

        for encoding in ["utf-8", "utf-16", "latin-1", "cp1252"]:
            try:
                with open(file_path, "r", encoding=encoding) as f:
                    text = f.read()
                    return text.strip() or "Empty text file"
            except UnicodeDecodeError:
                continue
            except Exception as e:
                logger.error(f"TXT read error: {e}")
        return "Error reading TXT file"

    def extract_text(self, file_path: str, extension: str) -> str:
        extension = extension.lower()
        if extension == ".pdf":
            return self.extract_text_from_pdf(file_path)
        if extension in [".doc", ".docx"]:
            return self.extract_text_from_docx(file_path)
        if extension == ".txt":
            return self.extract_text_from_txt(file_path)
        return f"Unsupported file format: {extension}"

    # --------------------------------------------------
    # SECTION EXTRACTION
    # --------------------------------------------------

    def extract_contact_info(self, text: str) -> str:
        if not text or text.startswith("Error"):
            return "Extraction failed"

        matches = []
        for pattern in self.contact_patterns:
            matches.extend(re.findall(pattern, text, re.IGNORECASE))
        return "; ".join(sorted(set(matches))) or "No contact information found"

    def extract_section(
        self,
        text: str,
        keywords: List[str],
        section_name: str,
        max_lines: int
    ) -> str:
        if not text or text.startswith("Error"):
            return f"No {section_name} section found"

        lines = [l.strip() for l in text.split("\n") if l.strip()]
        capture = False
        section_lines = []

        for line in lines:
            lower = line.lower()

            if any(k in lower for k in keywords):
                capture = True
                continue

            if capture and any(stop in lower for stop in self.stop_sections.get(section_name.lower(), [])):
                break

            if capture:
                section_lines.append(line)
                if len(section_lines) >= max_lines:
                    break

        return "\n".join(section_lines) or f"No {section_name} section found"

    def extract_experience(self, text: str) -> str:
        return self.extract_section(text, self.experience_keywords, "experience", 15)

    def extract_education(self, text: str) -> str:
        return self.extract_section(text, self.education_keywords, "education", 10)

    def extract_skills(self, text: str) -> str:
        return self.extract_section(text, self.skills_keywords, "skills", 8)

    # --------------------------------------------------
    # MAIN PARSER
    # --------------------------------------------------

    def parse_cv(self, file_path: str, extension: str) -> Dict:
        raw_text = self.extract_text(file_path, extension)

        if raw_text.startswith("Error"):
            return {
                "raw_text": raw_text,
                "contact_info": "Extraction failed",
                "experience": "Extraction failed",
                "education": "Extraction failed",
                "skills": "Extraction failed",
            }

        return {
            "raw_text": raw_text,
            "contact_info": self.extract_contact_info(raw_text),
            "experience": self.extract_experience(raw_text),
            "education": self.extract_education(raw_text),
            "skills": self.extract_skills(raw_text),
        }

    # --------------------------------------------------
    # MATCHING FUNCTION
    # --------------------------------------------------

    def match_with_criteria(
        self,
        parsed_cv: Dict,
        job_name: str = "",
        required_experience: Optional[int] = None,
        required_education: str = "",
        required_skills: Optional[List[str]] = None
    ) -> Dict:

        score = 0
        details = []

        if not parsed_cv:
            return {"score": 0, "details": ["Parsing failed"]}

        raw_text = parsed_cv.get("raw_text", "").lower()
        exp_text = parsed_cv.get("experience", "").lower()
        edu_text = parsed_cv.get("education", "").lower()
        skills_text = parsed_cv.get("skills", "").lower()

        required_skills = required_skills or []

        # -------- Job Name Matching --------
        if job_name:
            job_words = job_name.lower().split()
            if all(w in raw_text for w in job_words):
                score += 20
                details.append("Job title matched")
            else:
                details.append("Job title not matched")

        # -------- Experience Matching --------
        if required_experience:
            years = []
            duration_matches = re.findall(
                r'(\d+)\s*(?:\+)?(?:-|to)?\s*(\d+)?\s*(?:years|yrs|year)',
                exp_text
            )
            for y1, y2 in duration_matches:
                if y1.isdigit():
                    years.append(int(y1))
                if y2 and y2.isdigit():
                    years.append(int(y2))

            year_ranges = re.findall(
                r'\b((?:19|20)\d{2})\s*[-–]\s*((?:19|20)\d{2})\b',
                exp_text
            )
            for start, end in year_ranges:
                years.append(max(0, int(end) - int(start)))

            max_years = max(years) if years else 0
            if max_years >= required_experience:
                score += 20
                details.append(f"Experience OK ({max_years} yrs)")
            else:
                details.append(f"Experience insufficient ({max_years} yrs)")

        # -------- Education Matching --------
        if required_education:
            education_map = {
                "bachelor": ["bachelor", "bsc", "bs", "undergraduate"],
                "master": ["master", "msc", "ms", "graduate"],
                "phd": ["phd", "doctorate"]
            }

            req = required_education.lower()
            matched = any(
                any(term in edu_text for term in terms)
                for key, terms in education_map.items() if key == req
            )

            if matched:
                score += 20
                details.append("Education matched")
            else:
                details.append("Education not matched")

        # -------- Skills Matching --------
        if required_skills:
            cv_skills = re.split(r",|;|\n", skills_text)
            cv_skills = [s.strip().lower() for s in cv_skills if s.strip()]

            matched_skills = [
                s for s in required_skills
                if any(s.lower() in cv for cv in cv_skills)
            ]

            score += len(matched_skills) * 10
            if matched_skills:
                details.append(f"Skills matched: {', '.join(matched_skills)}")
            else:
                details.append("No skills matched")

        return {"score": score, "details": details}
